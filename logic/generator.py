from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# =====================================================
# CONSTANTS
# =====================================================
ROWS_MAIN = 17
COLS = 20
BOX = "□"

FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(10)

# =====================================================
# LOW-LEVEL HELPERS (XML SAFE)
# =====================================================
def merge_cells(table, r1, c1, r2, c2):
    table.cell(r1, c1).merge(table.cell(r2, c2))


def set_cell_vertical_text(cell):
    tc_pr = cell._tc.get_or_add_tcPr()

    for el in tc_pr.findall(qn("w:textDirection")):
        tc_pr.remove(el)

    text_dir = OxmlElement("w:textDirection")
    text_dir.set(qn("w:val"), "btLr")
    tc_pr.append(text_dir)


def set_cell_vertical_top(cell):
    tc_pr = cell._tc.get_or_add_tcPr()

    for el in tc_pr.findall(qn("w:vAlign")):
        tc_pr.remove(el)

    v_align = OxmlElement("w:vAlign")
    v_align.set(qn("w:val"), "top")
    tc_pr.append(v_align)


def set_cell_margins(cell, top=0, bottom=0, left=50, right=50):
    tc_pr = cell._tc.get_or_add_tcPr()

    for el in tc_pr.findall(qn("w:tcMar")):
        tc_pr.remove(el)

    tc_mar = OxmlElement("w:tcMar")
    for name, value in {
        "top": top,
        "bottom": bottom,
        "left": left,
        "right": right,
    }.items():
        node = OxmlElement(f"w:{name}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)

    tc_pr.append(tc_mar)


# =====================================================
# PARAGRAPH HELPERS
# =====================================================
def get_first_paragraph(cell):
    if cell.paragraphs:
        return cell.paragraphs[0]
    return cell.add_paragraph()


def style_run(run, bold=False):
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    run.bold = bold


def clear_first_paragraph(cell):
    p = cell.paragraphs[0]
    p.clear()
    p.paragraph_format.space_before = 0
    p.paragraph_format.space_after = 0


def add_simple(cell, text, *, center=False, bold=False, new_paragraph=False):
    if new_paragraph:
        p = cell.add_paragraph()
    else:
        p = get_first_paragraph(cell)

    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = 0
    p.paragraph_format.space_after = 0

    run = p.add_run(text)
    style_run(run, bold)


def add_label_value(cell, label, value="From Form"):
    p = cell.add_paragraph()
    p.paragraph_format.space_before = 0
    p.paragraph_format.space_after = 0

    hang = Cm(2.8)
    p.paragraph_format.first_line_indent = -hang
    p.paragraph_format.left_indent = hang

    r1 = p.add_run(f"{label}  ")
    style_run(r1, bold=False)

    r2 = p.add_run(value)
    style_run(r2, bold=True)


def add_center_form(cell):
    add_simple(cell, "From Form", center=True, bold=True, new_paragraph=True)


# =====================================================
# MAIN
# =====================================================
def generate_document(output_file: str, data: dict):
    doc = Document()
    section = doc.sections[0]

    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1)
    section.header_distance = Cm(1)
    section.footer_distance = Cm(1)

    # ================= TABLE 1 =================
    table1 = doc.add_table(rows=ROWS_MAIN, cols=COLS)
    table1.style = "Table Grid"
    table1.autofit = False

    MERGES = [
        (0, 0, 0, 19),
        (1, 0, 7, 0),
        (1, 1, 3, 9),
        (1, 10, 1, 17),
        (1, 18, 1, 19),
        (2, 10, 2, 19),
        (3, 10, 3, 19),
        (4, 1, 4, 9),
        (4, 10, 4, 19),
        (5, 1, 5, 4),
        (5, 5, 5, 9),
        (5, 10, 5, 14),
        (5, 15, 5, 19),
        (6, 1, 6, 9),
        (6, 10, 6, 19),
        (7, 1, 7, 9),
        (7, 10, 7, 19),
        (8, 0, 9, 9),
        (8, 10, 8, 19),
        (9, 10, 9, 19),
        (10, 0, 11, 9),
        (10, 10, 10, 19),
        (11, 10, 11, 19),
        (12, 0, 12, 9),
        (12, 10, 12, 19),
        (13, 0, 13, 9),
        (13, 10, 13, 19),
        (14, 0, 14, 19),
        (15, 0, 15, 9),
        (15, 10, 15, 19),
        (16, 0, 16, 19),
    ]

    for m in MERGES:
        merge_cells(table1, *m)

    # базова очистка + вирівнювання
    for row in table1.rows:
        for cell in row.cells:
            clear_first_paragraph(cell)
            set_cell_vertical_top(cell)
            set_cell_margins(cell, top=0, bottom=0)

    # ---------- TITLE ----------
    cell = table1.cell(0, 0)
    add_simple(
        cell,
        "Certificat de sănătate animală pentru circulația\n"
        "necomercială a păsărilor de companie",
        center=True,
        bold=True,
    )
    cell.paragraphs[0].runs[0].font.size = Pt(12)

    # ---------- PART I ----------
    cell = table1.cell(7, 0)
    set_cell_vertical_text(cell)
    add_simple(
        cell,
        "Partea I. Detalii privind transportul expediat",
        center=True,
        bold=True,
    )

    # ---------- I.1 ----------
    cell = table1.cell(3, 9)
    add_simple(cell, "I.1.   Expeditor")
    add_label_value(cell, "Nume")
    add_label_value(cell, "Adresă")
    add_label_value(cell, "Telefon")

    # ---------- I.2 ----------
    cell = table1.cell(1, 17)
    add_simple(cell, "I.2. Numărul de referință al certificatului")
    add_center_form(cell)

    # ---------- I.3 ----------
    cell = table1.cell(2, 19)
    add_simple(cell, "I.3. Autoritatea competentă centrală")
    add_simple(
        cell,
        "State servise of Ukraine on Food Safety and Consumer Protection (SSUFSCP)",
        center=True,
        bold=True,
        new_paragraph=True,
    )

    # ---------- I.4 ----------
    cell = table1.cell(3, 19)
    add_simple(cell, "I.4. Autoritatea competentă locală")
    add_simple(
        cell,
        "Lviv border inspection point (Lviv BIP)",
        center=True,
        bold=True,
        new_paragraph=True,
    )

    # ---------- I.5 ----------
    cell = table1.cell(4, 9)
    add_simple(cell, "I.5. Destinatar")
    for lbl in ("Nume", "Adresă", "Cod poștal", "Telefon"):
        add_label_value(cell, lbl)

    # ---------- I.6 ----------
    cell = table1.cell(4, 19)
    add_simple(cell, "I.6. Persoana responsabilă de lot")
    for lbl in ("Nume", "Adresă", "Cod poștal", "Telefon"):
        add_label_value(cell, lbl)

    # ---------- I.7–I.10 ----------
    add_simple(table1.cell(5, 4), "I.7. Țara de origine Cod ISO")
    add_simple(table1.cell(5, 4), "UA", center=True, bold=True, new_paragraph=True)

    add_simple(table1.cell(5, 9), "I.8. Regiunea de origine Cod")

    add_simple(table1.cell(5, 14), "I.9. Țara de destinație Cod ISO")
    add_simple(table1.cell(5, 14), "MD", center=True, bold=True, new_paragraph=True)

    add_simple(table1.cell(5, 19), "I.10.  Regiunea de destinație Cod")

    # ---------- I.11 ----------
    cell = table1.cell(6, 9)
    add_simple(cell, "I.11. Locul de origine")
    add_label_value(cell, "Nume")
    add_label_value(cell, "Adresă")

    # ---------- I.12 ----------
    cell = table1.cell(6, 19)
    add_simple(cell, "I.12.   Locul de destinație")
    add_label_value(cell, "Nume")
    add_label_value(cell, "Adresă")
    add_label_value(cell, "Cod poștal")

    # ---------- I.13 ----------
    add_simple(table1.cell(7, 9), "I.13.  Locul de încărcare")
    add_simple(table1.cell(7, 19), "I.13.  Locul de încărcare")

    # ---------- I.15 ----------
    cell = table1.cell(9, 9)
    add_simple(cell, "I.15.  Mijloace de transport")
    add_simple(cell, f"Avion {BOX}  Navă {BOX}  Vagon de cale ferată {BOX}", new_paragraph=True)
    add_simple(cell, f"Vehicul rutier {BOX}  Altele {BOX}", new_paragraph=True)
    add_label_value(cell, "Identificare:")

    # ---------- I.16 ----------
    cell = table1.cell(8, 19)
    add_simple(cell, "I.16. PIF de intrare în Republica Moldova")
    add_center_form(cell)

    # ---------- I.17 ----------
    cell = table1.cell(9, 19)
    add_simple(cell, "I.17. Numărul (numerele) CITES")
    add_center_form(cell)

    # ---------- I.18 ----------
    cell = table1.cell(11, 9)
    add_simple(cell, "I.18. Descrierea mărfii")
    add_center_form(cell)

    # ---------- I.19 ----------
    add_simple(
        table1.cell(10, 19),
        "I.19.  Codul mărfii (Codul SA):  From Form",
        bold=True,
    )

    # ---------- I.20 ----------
    add_simple(
        table1.cell(11, 19),
        "I.20. Cantitatea:  From Form",
        bold=True,
    )

    # ---------- I.21 ----------
    add_simple(table1.cell(12, 9), "I.21. Temperatura produselor")

    # ---------- I.22 ----------
    add_simple(table1.cell(12, 19), "I.22. Numărul de pachete")

    # ---------- I.23 ----------
    cell = table1.cell(13, 9)
    add_simple(cell, "I.23. Numărul sigiliului/containerului")
    add_center_form(cell)

    # ---------- I.24 ----------
    cell = table1.cell(13, 19)
    add_simple(cell, "I.24. Tipul de ambalaj")
    add_center_form(cell)

    # ---------- I.25 ----------
    cell = table1.cell(14, 9)
    add_simple(cell, "I.25. Mărfuri certificate pentru:")
    add_simple(
        cell,
        f"Animale de companie {BOX}      Carantină {BOX}",
        new_paragraph=True,
    )

    # ---------- I.26 ----------
    add_simple(table1.cell(15, 9), "I.26. Pentru tranzit")

    # ---------- I.27 ----------
    add_simple(table1.cell(15, 19), "I.27. Pentru import sau admitere în alte țări")

    # ---------- I.28 ----------
    cell = table1.cell(16, 19)
    add_simple(cell, "I.28. Identificarea mărfurilor")

    p = cell.add_paragraph()
    p.paragraph_format.space_before = 0
    p.paragraph_format.space_after = 0

    r = p.add_run("Specii (denumirea științifică):  ")
    style_run(r)
    r = p.add_run("From Form")
    style_run(r, bold=True)

    r = p.add_run("    Sistemul de identificare:  ")
    style_run(r)
    r = p.add_run("From Form")
    style_run(r, bold=True)

    p = cell.add_paragraph()
    p.paragraph_format.space_before = 0
    p.paragraph_format.space_after = 0

    r = p.add_run("Numărul de identificare:  ")
    style_run(r)
    r = p.add_run("From Form")
    style_run(r, bold=True)

    r = p.add_run("    Cantitatea:  ")
    style_run(r)
    r = p.add_run("From Form")
    style_run(r, bold=True)

    # ================= SAVE =================
    doc.save(output_file)
